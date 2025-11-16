# spec_ingestor.py
"""
Spec-Slice Ingestor

Parses a 3GPP-style DOCX and builds a Markdown context for a given keyword.
"""

import re
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

from config import DEFAULTS, IngestConfig
from models import Section, Reference
from utils import numeric_sort_key, render_table_markdown


class SpecSliceIngestor:
    """Parses DOCX and builds Markdown context per the specified rules."""

    def __init__(self, cfg: Optional[IngestConfig] = None):
        self.cfg = cfg or DEFAULTS.ingest
        self.root_sections: Dict[str, Section] = {}
        self.current_section: Optional[Section] = None
        # Enable parsing immediately if no start_section; otherwise wait until we hit start >=
        self.parsing_enabled = self.cfg.start_section is None
        self.section_pattern = re.compile(self.cfg.section_header_regex)
        self.reference_patterns = self._build_reference_patterns()

    # ---------- Public API ----------

    def parse(self, docx_path: str) -> Dict[str, Section]:
        """Parse the DOCX and populate the internal section tree."""
        doc = Document(docx_path)
        for element in doc.element.body:
            if element.tag.endswith("p"):
                self._process_paragraph(Paragraph(element, doc))
            elif element.tag.endswith("tbl") and self.cfg.include_tables:
                self._process_table(Table(element, doc))
        return self.root_sections

    def build_context_markdown_for_keyword(
        self,
        keyword: str,
        max_title_hits: int = 50,
        max_non_title_hits: int = 50,
        *,
        resolve_local_refs: bool = True,
        recursive_ref_resolution: bool = True,
    ) -> str:
        """Build a Markdown context for the given keyword."""
        kw = keyword.lower()

        # Helper functions
        def _is_ancestor(anc: str, desc: str) -> bool:
            return desc.startswith(anc + ".")

        def _prune_title_hits(hits: List[Section]) -> List[Section]:
            hits_sorted = sorted(hits, key=lambda s: numeric_sort_key(s.number))
            kept: List[Section] = []
            for sec in hits_sorted:
                if any(_is_ancestor(k.number, sec.number) for k in kept):
                    continue
                kept.append(sec)
            return kept

        # Global cache for resolved clause expansions
        resolved_cache: Dict[str, str] = {}
        resolving_stack: set[str] = set()

        def _collect_local_clause_refs_from_table(tbl: dict) -> List[str]:
            out: List[str] = []
            for r in tbl.get("references", []):
                if r.get("ref_type") == "clause":
                    clause = (r.get("clause") or "").strip()
                    doc = (r.get("document") or "").strip()
                    if clause and (not doc or doc == "-" or "TS" not in doc):
                        out.append(clause)
            return out

        def _collect_local_clause_refs_from_section(sec: Section) -> List[str]:
            out: List[str] = []
            for ref in sec.references:
                if ref.ref_type == "clause" and ref.clause:
                    if not ref.document or ref.document.strip() in {"", "-"}:
                        out.append(ref.clause.strip())
            for tbl in sec.tables:
                out.extend(_collect_local_clause_refs_from_table(tbl))
            # uniq, preserve order
            seen = set()
            uniq = []
            for c in out:
                if c not in seen:
                    seen.add(c)
                    uniq.append(c)
            return uniq

        def _render_clause_body(sec: Section) -> str:
            parts: List[str] = [f"#### §{sec.number} {sec.title}\n"]
            for t in sec.text:
                parts.append(t)
                parts.append("")
            for tbl in sec.tables:
                tmd = render_table_md(tbl, visited_for_this_chain=set())
                if tmd:
                    parts.append(tmd)
                    parts.append("")
            return "\n".join(parts).rstrip()

        def _resolve_clause_recursive(clause_num: str, visited_for_this_chain: set[str]) -> str:
            if clause_num in resolving_stack:
                return f"_Circular reference detected; already in stack: §{clause_num}_"

            if clause_num in resolved_cache:
                sec = self._find_by_number(clause_num)
                title = f" {sec.title}" if sec else ""
                return f"_See earlier expansion of §{clause_num}{title}._"

            sec = self._find_by_number(clause_num)
            if not sec:
                return f"_Could not resolve local clause §{clause_num}_"

            resolving_stack.add(clause_num)
            visited_for_this_chain.add(clause_num)

            md = _render_clause_body(sec)

            if recursive_ref_resolution:
                child_refs = _collect_local_clause_refs_from_section(sec)
                child_blocks = []
                for cnum in child_refs:
                    if cnum in visited_for_this_chain:
                        continue
                    child_block = _resolve_clause_recursive(cnum, visited_for_this_chain)
                    if child_block:
                        child_blocks.append(child_block)
                if child_blocks:
                    md = md + "\n\n" + "\n\n".join(child_blocks)

            resolving_stack.remove(clause_num)
            resolved_cache[clause_num] = md
            return md

        def render_table_md(tbl: dict, *, visited_for_this_chain: set[str]) -> str:
            headers = tbl.get('headers', [])
            rows = tbl.get('rows', [])
            if not headers and not rows:
                return ""
            md = ""
            if headers:
                md += "| " + " | ".join(h or "" for h in headers) + " |\n"
                md += "| " + " | ".join(["---"] * len(headers)) + " |\n"
            for row in rows:
                data = row[:len(headers)] + [""] * (len(headers) - len(row)) if headers else row
                md += "| " + " | ".join(str(c or "") for c in data) + " |\n"

            refs = tbl.get("references", [])
            if refs:
                seen_keys = set()
                uniq_refs = []
                for r in refs:
                    key = (r.get("ref_type"), (r.get("clause") or "").strip(), (r.get("document") or "").strip())
                    if key in seen_keys:
                        continue
                    seen_keys.add(key)
                    uniq_refs.append(r)

                if uniq_refs:
                    md += "\n_References under this table:_\n"
                    for i, r in enumerate(uniq_refs, 1):
                        rtype = "Clause" if r.get("ref_type") == "clause" else "Document"
                        clause = (r.get('clause') or "").strip()
                        clause_md = f"`{clause}`" if clause else "`-`"
                        doc = (r.get('document') or "").strip()
                        doc_md = f"`{doc}`" if doc else "`-`"
                        ctx = (r.get("context") or "").strip().replace("\n", " ")

                        md += f"  {i}. **{rtype}** | clause: {clause_md} | document: {doc_md}\n"
                        if ctx:
                            md += f"     - _\"{ctx}\"_\n"

                        if resolve_local_refs and r.get("ref_type") == "clause" and (not doc or doc in {"", "-"}):
                            if clause:
                                expanded = _resolve_clause_recursive(clause, visited_for_this_chain)
                                if expanded:
                                    md += "     - **Resolved (local, full clause)**:\n"
                                    md += "      " + expanded.replace("\n", "\n      ") + "\n"
            return md

        def render_section_md(sec: Section, include_subtree: bool = True) -> str:
            parts = [f"### §{sec.number} {sec.title}\n"]
            for t in sec.text:
                parts.append(t)
                parts.append("")
            for tbl in sec.tables:
                tmd = render_table_md(tbl, visited_for_this_chain=set())
                if tmd:
                    parts.append(tmd)
                    parts.append("")
            if include_subtree and sec.subsections:
                for k in sorted(sec.subsections.keys(), key=numeric_sort_key):
                    parts.append(render_section_md(sec.subsections[k], True))
            return "\n".join(parts)

        # Gather sections
        all_sections = self._iter_sections()

        # Title hits
        title_hits: List[Section] = [s for s in all_sections if kw in s.title.lower()]
        title_hits.sort(key=lambda s: "\n".join([s.title] + s.text).lower().count(kw), reverse=True)
        title_hits = title_hits[:max_title_hits] if max_title_hits else title_hits
        title_hits = _prune_title_hits(title_hits)

        # Build included numbers set
        included_numbers: set[str] = set()
        def collect_subtree_numbers(s: Section):
            included_numbers.add(s.number)
            for sub in s.subsections.values():
                collect_subtree_numbers(sub)
        for th in title_hits:
            collect_subtree_numbers(th)

        # Non-title paragraph hits
        para_hits: List[tuple[Section, str]] = []
        for s in all_sections:
            if s.number in included_numbers:
                continue
            for para in s.text:
                if kw in para.lower():
                    para_hits.append((s, para))
        para_hits.sort(key=lambda p: p[1].lower().count(kw), reverse=True)
        para_hits = para_hits[:max_non_title_hits] if max_non_title_hits else para_hits

        # Render output
        out: List[str] = []
        out.append(f"# Context for `{keyword}`\n")

        out.append("## Matching Sections (title hits)\n")
        if title_hits:
            for s in sorted(title_hits, key=lambda sec: numeric_sort_key(sec.number)):
                out.append(render_section_md(s, include_subtree=True))
                out.append("")
        else:
            out.append("_None_\n")

        if para_hits:
            out.append("## Matching Paragraphs (non-title hits)\n")
            by_sec: Dict[str, List[str]] = {}
            sec_map: Dict[str, Section] = {}
            for s, p in para_hits:
                by_sec.setdefault(s.number, []).append(p)
                sec_map[s.number] = s
            for secnum in sorted(by_sec.keys(), key=numeric_sort_key):
                s = sec_map[secnum]
                out.append(f"### From §{s.number} {s.title}\n")
                for para in by_sec[secnum]:
                    out.append(para)
                    out.append("")

        return "\n".join(out).rstrip() + "\n"

    # ---------- Internal methods ----------

    def _find_matches(self, keyword: str) -> Tuple[List[Section], List[Tuple[Section, str]]]:
        """Return (full_sections, paragraph_hits) in document order."""
        kw = keyword.lower()
        full_sections: List[Section] = []
        paragraph_hits: List[Tuple[Section, str]] = []

        for sec in self._iter_sections():
            if kw in sec.title.lower():
                full_sections.append(sec)
                continue
            for para in sec.text:
                if kw in para.lower():
                    paragraph_hits.append((sec, para))
        return full_sections, paragraph_hits

    def _render_section_full_md(self, sec: Section, *, include_subtree: bool = False, level: int = 3) -> str:
        hlevel = max(1, min(level, 6))
        parts: List[str] = [f"{'#' * hlevel} §{sec.number} {sec.title}", ""]

        for t in sec.text:
            parts.append(t)
            parts.append("")

        for idx, tbl in enumerate(sec.tables, 1):
            parts.append(f"**Table {idx}**")
            parts.append("")
            table_md = render_table_markdown(tbl, include_references=True)
            parts.append(table_md or "_<empty table>_")
            parts.append("")

        if include_subtree and sec.subsections:
            for sub in sec.subsections.values():
                parts.append(self._render_section_full_md(sub, include_subtree=True, level=hlevel + 1))
                parts.append("")

        return "\n".join(parts).rstrip()

    def _build_reference_patterns(self):
        return {
            "document_with_clause": [
                re.compile(r"\b(?:clause|subclause|section)\s+(\d+(?:\.\d+)*)\s+of\s+((?:3GPP\s+)?TS\s+\d+\.\d+(?:\s*\[\d+\])?)", re.IGNORECASE),
                re.compile(r"\b(\d+(?:\.\d+)*)\s+of\s+((?:3GPP\s+)?TS\s+\d+\.\d+(?:\s*\[\d+\])?)", re.IGNORECASE),
            ],
            "document": [
                re.compile(r"\b((?:3GPP\s+)?TS\s+\d+\.\d+(?:\s*\[\d+\])?)", re.IGNORECASE),
                re.compile(r"\b(RFC\s+\d+)", re.IGNORECASE),
                re.compile(r"\b(ITU-T\s+[A-Z]\.\d+)", re.IGNORECASE),
            ],
            "clause": [
                re.compile(r"\b(?:clause|subclause|section)\s+(\d+(?:\.\d+)*)", re.IGNORECASE),
                re.compile(r"\b(\d+(?:\.\d+){2,})\b"),
            ],
        }

    def _process_paragraph(self, paragraph: Paragraph):
        text = paragraph.text.strip()
        if not text:
            return

        section_match = self.section_pattern.match(text)
        if section_match and self._is_section_header(paragraph):
            section_number = section_match.group(1)
            section_title = section_match.group(2).strip()
            if self._should_parse_section(section_number, section_title):
                self.parsing_enabled = True
                self._create_section(section_number, section_title, text)
            else:
                self._update_parsing_state(section_number)
        else:
            if self.parsing_enabled and self.current_section and not self._is_figure_caption(text):
                self.current_section.text.append(text)
                refs = self._extract_references(text, context=text)
                self.current_section.references.extend(refs)

    def _process_table(self, table: Table):
        if not self.parsing_enabled or not self.current_section:
            return

        table_data = {"headers": [], "rows": [], "references": []}
        if table.rows:
            header_row = table.rows[0]
            table_data["headers"] = [cell.text.strip() for cell in header_row.cells]

            for row in table.rows[1:]:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data["rows"].append(row_data)

                row_context = "  ".join(row_data)
                refs = self._extract_references(" ".join(row_data), context=row_context)
                for ref in refs:
                    table_data["references"].append({
                        "ref_type": ref.ref_type,
                        "clause": ref.clause,
                        "document": ref.document,
                        "context": ref.context,
                    })

        self.current_section.tables.append(table_data)

    def _should_parse_section(self, number: str, title: str) -> bool:
        if self.cfg.exclude_toc and ("contents" in title.lower() or "table of contents" in title.lower()):
            return False
        if self.cfg.exclude_appendix and (title.lower().startswith("appendix") or number.upper().startswith("A")):
            return False

        if self.cfg.start_section or self.cfg.end_section:
            try:
                if self.cfg.start_section and self._compare_sections(number, self.cfg.start_section) < 0:
                    return False
                if self.cfg.end_section and self._compare_sections(number, self.cfg.end_section) > 0:
                    return False
            except (ValueError, IndexError):
                if number and number[0].isalpha():
                    return not self.cfg.exclude_appendix
        return True

    def _compare_sections(self, s1: str, s2: str) -> int:
        def parse(x: str) -> List[int]:
            return [int(t) for t in x.split(".")]
        a, b = parse(s1), parse(s2)
        for i in range(max(len(a), len(b))):
            v1 = a[i] if i < len(a) else 0
            v2 = b[i] if i < len(b) else 0
            if v1 < v2:
                return -1
            if v1 > v2:
                return 1
        return 0

    def _update_parsing_state(self, number: str):
        try:
            if self.cfg.start_section and self._compare_sections(number, self.cfg.start_section) < 0:
                self.parsing_enabled = False
            elif self.cfg.end_section and self._compare_sections(number, self.cfg.end_section) > 0:
                self.parsing_enabled = False
        except (ValueError, IndexError):
            pass

    def _create_section(self, number: str, title: str, full_title: str):
        level = len(number.split("."))
        s = Section(number=number, title=title, full_title=full_title, level=level)
        parent = self._find_parent(number)
        if parent:
            s.parent = parent
            parent.subsections[number] = s
        else:
            self.root_sections[number] = s
        self.current_section = s

    def _find_parent(self, number: str) -> Optional[Section]:
        parts = number.split(".")
        if len(parts) <= 1:
            return None
        parent_number = ".".join(parts[:-1])
        return self._find_by_number(parent_number)

    def _find_by_number(self, number: str) -> Optional[Section]:
        def search(d: Dict[str, Section]) -> Optional[Section]:
            if number in d:
                return d[number]
            for sec in d.values():
                r = search(sec.subsections)
                if r:
                    return r
            return None
        return search(self.root_sections)

    def _iter_sections(self) -> List[Section]:
        out: List[Section] = []
        def rec(d: Dict[str, Section]):
            for s in d.values():
                out.append(s)
                rec(s.subsections)
        rec(self.root_sections)
        return out

    def _is_section_header(self, p: Paragraph) -> bool:
        if p.style and p.style.name and p.style.name.startswith("Heading"):
            return True
        if p.runs:
            bold_chars = 0
            total_chars = 0
            for run in p.runs:
                if run.text.strip():
                    total_chars += len(run.text)
                    if getattr(run, "bold", False):
                        bold_chars += len(run.text)
            if total_chars > 0 and bold_chars / total_chars > 0.5:
                return True
        return False

    def _is_figure_caption(self, text: str) -> bool:
        for pat in [r"^Figure\s+\d+", r"^Table\s+\d+", r"Figure\s+\d+\.\d+"]:
            if re.match(pat, text, re.IGNORECASE):
                return True
        return False

    def _extract_references(self, text: str, context: Optional[str] = None) -> List[Reference]:
        context = context or text
        refs: List[Reference] = []
        consumed: set[int] = set()

        # Document with clause
        for pat in self.reference_patterns["document_with_clause"]:
            for m in pat.finditer(text):
                refs.append(Reference("clause", m.group(1), m.group(2), context))
                consumed.update(range(m.start(), m.end()))

        # Standalone documents
        for pat in self.reference_patterns["document"]:
            for m in pat.finditer(text):
                if any(pos in consumed for pos in range(m.start(), m.end())):
                    continue
                refs.append(Reference("document", None, m.group(1), context))
                consumed.update(range(m.start(), m.end()))

        # Standalone clauses
        for pat in self.reference_patterns["clause"]:
            for m in pat.finditer(text):
                if any(pos in consumed for pos in range(m.start(), m.end())):
                    continue
                refs.append(Reference("clause", m.group(1), None, context))

        return refs


if __name__ == "__main__":
    import argparse, sys
    ap = argparse.ArgumentParser(description="Ingest DOCX and build context Markdown for a keyword")
    ap.add_argument("docx", help="Path to DOCX")
    ap.add_argument("keyword", help="Keyword / procedure name")
    ap.add_argument("--start", default=DEFAULTS.ingest.start_section)
    ap.add_argument("--end", default=DEFAULTS.ingest.end_section)
    args = ap.parse_args()

    ing = SpecSliceIngestor(IngestConfig(start_section=args.start, end_section=args.end))
    ing.parse(args.docx)

    md = ing.build_context_markdown_for_keyword(args.keyword)
    sys.stdout.write(md)
